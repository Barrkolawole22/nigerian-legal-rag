import io

from fastapi import APIRouter, UploadFile, File, HTTPException
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.services.vectorstore import create_session_store
from app.models.schemas import UploadResponse
from app.config import get_settings

router = APIRouter(prefix="/documents", tags=["documents"])
settings = get_settings()

SUPPORTED = {".pdf", ".docx"}


def _pdf_to_docs(content: bytes, filename: str) -> list[Document]:
    reader = PdfReader(io.BytesIO(content))
    return [
        Document(
            page_content=page.extract_text() or "",
            metadata={"source": filename, "page": i},
        )
        for i, page in enumerate(reader.pages)
        if (page.extract_text() or "").strip()
    ]


def _docx_to_docs(content: bytes, filename: str) -> list[Document]:
    doc = DocxDocument(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": filename, "page": 0})]


@router.post("/upload", response_model=UploadResponse)
async def upload(file: UploadFile = File(...)):
    name = file.filename or "upload"
    ext = ("." + name.rsplit(".", 1)[-1].lower()) if "." in name else ""

    if ext not in SUPPORTED:
        raise HTTPException(status_code=400, detail=f"Unsupported type. Accepted: {SUPPORTED}")

    content = await file.read()

    if len(content) > settings.MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"File exceeds {settings.MAX_UPLOAD_MB}MB")

    raw_docs = _pdf_to_docs(content, name) if ext == ".pdf" else _docx_to_docs(content, name)

    if not raw_docs:
        raise HTTPException(status_code=422, detail="No text could be extracted")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(raw_docs)

    if not chunks:
        raise HTTPException(status_code=422, detail="Document produced no usable chunks")

    session_id = create_session_store(chunks)

    return UploadResponse(session_id=session_id, filename=name, chunks=len(chunks))
